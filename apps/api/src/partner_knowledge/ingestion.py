"""Build the persistent local Partner knowledge index from approved sources."""

import csv
import hashlib
import json
import logging
import sqlite3
from collections.abc import Callable, Iterable, Mapping
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from uuid import uuid4
from zipfile import BadZipFile

import chromadb
from chromadb.errors import ChromaError
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from src.partner_knowledge.cohere_embeddings import (
    CohereEmbeddingClient,
    CohereEmbeddingConfigurationError,
    CohereEmbeddingError,
)
from src.partner_knowledge.config import (
    PartnerKnowledgeSettings,
    get_partner_knowledge_settings,
)
from src.partner_knowledge.embedding_budget import (
    EmbeddingUsageLedger,
    EmbeddingUsageLedgerError,
)
from src.partner_knowledge.index_storage import (
    ACTIVATION_LOCK_NAME,
    COLLECTION_NAME,
    INDEX_LOCK_NAME,
    STAGING_COLLECTION_PREFIX,
    PartnerKnowledgeIndexLockError,
    activate_collection,
    active_collection_name,
    partner_knowledge_index_lock,
)
from src.provider_errors import ProviderRateLimitError

_SOURCE_FINGERPRINT_KEY = "source_fingerprint"
logger = logging.getLogger(__name__)
_APPROVED_SOURCES = {
    "Catálogo de Produtos e Ingredientes - Café Aurora.pdf": (
        "Catálogo de Produtos e Ingredientes - Café Aurora",
        "pdf",
    ),
    "CA-COM-PLA-001_Controle_de_Estoque.csv": ("Controle de Estoque", "csv"),
    "CA-TEC-CAD-001_Configuracao_das_Unidades.json": (
        "Configuração das Unidades",
        "json",
    ),
    "Manual de Operações das Unidades - Café Aurora.pdf": (
        "Manual de Operações das Unidades - Café Aurora",
        "pdf",
    ),
    "CA-QUA-GUI-001_Guia_de_Atendimento_ao_Cliente.docx": (
        "Guia de Atendimento ao Cliente",
        "docx",
    ),
    "CA-FIN-POL-001_Politica_de_Despesas_e_Reembolsos.docx": (
        "Política de Despesas e Reembolsos",
        "docx",
    ),
}


class PartnerKnowledgeIngestionError(RuntimeError):
    """Raised when approved Partner knowledge cannot be safely indexed."""


@dataclass(frozen=True)
class IngestionResult:
    indexed_chunks: int
    index_path: Path
    skipped: bool = False


@dataclass(frozen=True)
class _Chunk:
    text: str
    document_name: str
    location: str
    technical_location: str
    source_kind: str

    @property
    def id(self) -> str:
        identity = "\n".join((self.document_name, self.technical_location, self.text))
        return hashlib.sha256(identity.encode()).hexdigest()

    @property
    def metadata(self) -> dict[str, str]:
        return {
            "document_name": self.document_name,
            "location": self.location,
            "technical_location": self.technical_location,
            "source_kind": self.source_kind,
        }


class PartnerKnowledgeIngestor:
    """Explicit, idempotent builder for the approved Partner knowledge index."""

    def __init__(
        self,
        settings: PartnerKnowledgeSettings,
        *,
        embed_documents: Callable[[list[str]], list[list[float]]],
        embedding_metadata: Mapping[str, str] | None = None,
    ):
        self._settings = settings
        self._embed_documents = embed_documents
        self._embedding_metadata = dict(embedding_metadata or {})

    def ingest(self) -> IngestionResult:
        try:
            self._settings.partner_index_path.mkdir(parents=True, exist_ok=True)
            with partner_knowledge_index_lock(
                self._settings.partner_index_path,
                exclusive=True,
                non_blocking=True,
                lock_name=INDEX_LOCK_NAME,
                lock_directory=self._settings.runtime_data_path,
            ):
                return self._ingest_locked()
        except BlockingIOError as exc:
            raise PartnerKnowledgeIngestionError(
                "Another Partner knowledge ingestion operation is already running."
            ) from exc
        except PartnerKnowledgeIndexLockError as exc:
            raise PartnerKnowledgeIngestionError(
                "Unable to open the Partner knowledge ingestion lock at "
                f"{self._settings.partner_index_path}."
            ) from exc
        except OSError as exc:
            raise PartnerKnowledgeIngestionError(
                "Unable to access the Partner knowledge index at "
                f"{self._settings.partner_index_path}."
            ) from exc

    def _ingest_locked(self) -> IngestionResult:
        chunks = list(self._load_approved_chunks())
        if not chunks:
            raise PartnerKnowledgeIngestionError(
                "No approved Partner knowledge documents were found in "
                f"{self._settings.partner_document_source}."
            )
        fingerprint = _source_fingerprint(chunks, self._embedding_metadata)
        try:
            client = chromadb.PersistentClient(
                path=str(self._settings.partner_index_path)
            )
        except (ChromaError, OSError, ValueError, sqlite3.Error) as exc:
            raise PartnerKnowledgeIngestionError(
                "Unable to open the Partner knowledge index at "
                f"{self._settings.partner_index_path}."
            ) from exc
        expected_metadata = {
            **self._embedding_metadata,
            _SOURCE_FINGERPRINT_KEY: fingerprint,
        }
        if self._embedding_metadata:
            try:
                index_is_current = _active_index_is_current(
                    client,
                    index_path=self._settings.partner_index_path,
                    expected_metadata=expected_metadata,
                    expected_chunk_count=len(chunks),
                )
            except (OSError, ValueError, sqlite3.Error) as exc:
                raise PartnerKnowledgeIngestionError(
                    "The active Partner knowledge index pointer is unreadable at "
                    f"{self._settings.partner_index_path}."
                ) from exc
            if index_is_current:
                return IngestionResult(
                    indexed_chunks=len(chunks),
                    index_path=self._settings.partner_index_path,
                    skipped=True,
                )
        try:
            embeddings = self._embed_documents([chunk.text for chunk in chunks])
        except (
            EmbeddingUsageLedgerError,
            ProviderRateLimitError,
        ) as exc:
            raise PartnerKnowledgeIngestionError(str(exc)) from exc
        except CohereEmbeddingConfigurationError as exc:
            raise PartnerKnowledgeIngestionError(str(exc)) from exc
        except CohereEmbeddingError as exc:
            raise PartnerKnowledgeIngestionError(
                "Unable to create embeddings for Partner knowledge."
            ) from exc
        except Exception as exc:
            raise PartnerKnowledgeIngestionError(
                "Unable to create embeddings for Partner knowledge."
            ) from exc
        if len(embeddings) != len(chunks):
            raise PartnerKnowledgeIngestionError(
                "Embedding provider returned a different number of vectors "
                "than documents."
            )

        staging_name = f"{STAGING_COLLECTION_PREFIX}{uuid4().hex}"
        collection_metadata = {
            "hnsw:space": "cosine",
            **expected_metadata,
        }
        try:
            collection = client.create_collection(
                staging_name,
                embedding_function=None,
                metadata=collection_metadata,
            )
            collection.add(
                ids=[chunk.id for chunk in chunks],
                documents=[chunk.text for chunk in chunks],
                metadatas=[chunk.metadata for chunk in chunks],
                embeddings=embeddings,
            )
            if collection.count() != len(chunks):
                raise ValueError(
                    "Staged collection contains an unexpected record count."
                )
        except (ChromaError, OSError, ValueError, sqlite3.Error) as exc:
            try:
                client.delete_collection(staging_name)
            except (ChromaError, OSError, sqlite3.Error):
                pass
            raise PartnerKnowledgeIngestionError(
                "Unable to write the Partner knowledge index at "
                f"{self._settings.partner_index_path}."
            ) from exc
        activation_committed = False
        try:
            with partner_knowledge_index_lock(
                self._settings.partner_index_path,
                exclusive=True,
                lock_name=ACTIVATION_LOCK_NAME,
                lock_directory=self._settings.runtime_data_path,
            ):
                activate_collection(self._settings.partner_index_path, staging_name)
                activation_committed = True
                _garbage_collect_collections(
                    client, active_collection_name=staging_name
                )
        except (OSError, ValueError, sqlite3.Error) as exc:
            if activation_committed:
                logger.warning(
                    "Partner knowledge activation committed with a cleanup or "
                    "lock-finalization issue",
                    extra={"extra_data": {"error": str(exc)}},
                )
            else:
                try:
                    client.delete_collection(staging_name)
                except (ChromaError, OSError, sqlite3.Error):
                    pass
                raise PartnerKnowledgeIngestionError(
                    "Unable to activate the rebuilt Partner knowledge index at "
                    f"{self._settings.partner_index_path}."
                ) from exc
        return IngestionResult(
            indexed_chunks=len(chunks), index_path=self._settings.partner_index_path
        )

    def _load_approved_chunks(self) -> Iterable[_Chunk]:
        source_root = self._settings.partner_document_source
        if not source_root.is_dir():
            raise PartnerKnowledgeIngestionError(
                f"Partner document source is not a directory: {source_root}"
            )
        for filename, (document_name, source_kind) in _APPROVED_SOURCES.items():
            path = source_root / filename
            if not path.is_file():
                continue
            if source_kind == "pdf":
                yield from self._extract_pdf(path, document_name)
            elif source_kind == "docx":
                yield from self._extract_docx(path, document_name)
            elif source_kind == "csv":
                yield from self._extract_csv(path, document_name)
            else:
                yield from self._extract_json(path, document_name)

    @staticmethod
    def _extract_pdf(path: Path, document_name: str) -> Iterable[_Chunk]:
        try:
            pages = PdfReader(path).pages
            extracted_pages = [page.extract_text().strip() for page in pages]
        except (PdfReadError, OSError) as exc:
            raise PartnerKnowledgeIngestionError(
                f"Approved document {path.name} has no readable native text; "
                "OCR is not supported."
            ) from exc
        if not any(extracted_pages):
            raise PartnerKnowledgeIngestionError(
                f"Approved document {path.name} has no readable native text; "
                "OCR is not supported."
            )
        for page_number, text in enumerate(extracted_pages, start=1):
            if text:
                yield _Chunk(
                    text,
                    document_name,
                    f"Página {page_number}",
                    f"page:{page_number}",
                    "pdf",
                )

    @staticmethod
    def _extract_docx(path: Path, document_name: str) -> Iterable[_Chunk]:
        try:
            document = Document(path)
        except (BadZipFile, PackageNotFoundError, OSError) as exc:
            raise PartnerKnowledgeIngestionError(
                f"Approved document {path.name} has no readable native text; "
                "OCR is not supported."
            ) from exc
        paragraphs = document.paragraphs
        sections: list[tuple[str, list[str]]] = []
        heading, body = "Introdução", []
        for paragraph in paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style.name.startswith("Heading"):
                if body:
                    sections.append((heading, body))
                heading, body = text, []
            else:
                body.append(text)
        if body:
            sections.append((heading, body))
        if not sections:
            raise PartnerKnowledgeIngestionError(
                f"Approved document {path.name} has no readable native text; "
                "OCR is not supported."
            )
        for index, (heading, body) in enumerate(sections, start=1):
            yield _Chunk(
                "\n".join([heading, *body]),
                document_name,
                f"Seção: {heading}",
                f"section:{index}",
                "docx",
            )
        for table_number, table in enumerate(document.tables, start=1):
            for row_number, row in enumerate(table.rows, start=1):
                cells = [cell.text.strip() for cell in row.cells]
                if not any(cells):
                    continue
                yield _Chunk(
                    " | ".join(cells),
                    document_name,
                    f"Tabela {table_number}, linha {row_number}",
                    f"table:{table_number}:row:{row_number}",
                    "docx",
                )

    @staticmethod
    def _extract_csv(path: Path, document_name: str) -> Iterable[_Chunk]:
        try:
            with path.open(encoding="utf-8-sig", newline="") as source_file:
                for row_number, row in enumerate(csv.DictReader(source_file), start=2):
                    values = {key: value for key, value in row.items() if value}
                    if not values:
                        continue
                    unit = values.get("Código da unidade", "unidade não identificada")
                    item = values.get("Código do item", "item não identificado")
                    description = values.get("Descrição", "")
                    location = f"Unidade {unit}, item {item}" + (
                        f" - {description}" if description else ""
                    )
                    yield _Chunk(
                        "\n".join(f"{key}: {value}" for key, value in values.items()),
                        document_name,
                        location,
                        f"row:{row_number}",
                        "csv",
                    )
        except (OSError, UnicodeDecodeError, csv.Error) as exc:
            raise PartnerKnowledgeIngestionError(
                f"Approved document {path.name} could not be read as CSV."
            ) from exc

    @staticmethod
    def _extract_json(path: Path, document_name: str) -> Iterable[_Chunk]:
        try:
            content: dict[str, Any] = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, UnicodeDecodeError, json.JSONDecodeError) as exc:
            raise PartnerKnowledgeIngestionError(
                f"Approved document {path.name} could not be read as JSON."
            ) from exc
        global_configuration = {
            key: value
            for key, value in content.items()
            if key not in {"unidades", "documento"}
        }
        if global_configuration:
            yield _Chunk(
                json.dumps(global_configuration, ensure_ascii=False, indent=2),
                document_name,
                "Configurações globais",
                "json:global",
                "json",
            )
        for unit_index, unit in enumerate(content.get("unidades", [])):
            unit_code, unit_name = (
                unit.get("codigo", "unidade não identificada"),
                unit.get("nome", ""),
            )
            location = f"Unidade {unit_code}" + (f" - {unit_name}" if unit_name else "")
            yield _Chunk(
                json.dumps(unit, ensure_ascii=False, indent=2),
                document_name,
                location,
                f"json:unidades[{unit_index}]",
                "json",
            )


def _source_fingerprint(
    chunks: list[_Chunk], embedding_metadata: Mapping[str, str]
) -> str:
    payload = {
        "chunk_ids": [chunk.id for chunk in chunks],
        "embedding_metadata": dict(sorted(embedding_metadata.items())),
    }
    encoded = json.dumps(payload, ensure_ascii=False, sort_keys=True).encode()
    return hashlib.sha256(encoded).hexdigest()


def _active_index_is_current(
    client: Any,
    *,
    index_path: Path,
    expected_metadata: Mapping[str, str],
    expected_chunk_count: int,
) -> bool:
    collection_name = active_collection_name(index_path)
    try:
        collection = client.get_collection(collection_name)
        if collection.count() != expected_chunk_count:
            return False
        metadata = collection.metadata or {}
    except (ChromaError, OSError, ValueError, sqlite3.Error):
        return False
    return all(
        str(metadata.get(key)) == value for key, value in expected_metadata.items()
    )


def _garbage_collect_collections(client: Any, *, active_collection_name: str) -> None:
    try:
        collections = client.list_collections()
    except (ChromaError, OSError, ValueError, sqlite3.Error) as exc:
        logger.warning(
            "Unable to inspect stale Partner knowledge collections after activation",
            extra={"extra_data": {"error": str(exc)}},
        )
        return
    for collection in collections:
        collection_name = collection.name
        if collection_name == active_collection_name or not (
            collection_name == COLLECTION_NAME
            or collection_name.startswith(STAGING_COLLECTION_PREFIX)
        ):
            continue
        try:
            client.delete_collection(collection_name)
        except (ChromaError, OSError, ValueError, sqlite3.Error) as exc:
            logger.warning(
                "Unable to remove stale Partner knowledge collection",
                extra={
                    "extra_data": {
                        "collection_name": collection_name,
                        "error": str(exc),
                    }
                },
            )


def main() -> None:
    """Run explicit local ingestion with Cohere document embeddings."""
    settings = get_partner_knowledge_settings()
    ledger = EmbeddingUsageLedger(
        settings.embedding_usage_ledger_path,
        monthly_limit=settings.embedding_monthly_call_limit,
    )
    embeddings = CohereEmbeddingClient(
        settings.cohere_api_key,
        model=settings.embedding_model,
        embedding_dimension=settings.embedding_dimension,
        ledger=ledger,
    )
    result = PartnerKnowledgeIngestor(
        settings,
        embed_documents=embeddings.embed_documents,
        embedding_metadata=embeddings.collection_metadata,
    ).ingest()
    usage = embeddings.usage_snapshot()
    operation = "Index already current" if result.skipped else "Indexed"
    projected_calls = (
        0
        if result.skipped
        else embeddings.projected_document_calls(result.indexed_chunks)
    )
    print(
        f"{operation} {result.indexed_chunks} Partner knowledge chunks at "
        f"{result.index_path}. Projected Embed calls: {projected_calls}. "
        f"UTC-month usage: {usage.used_calls}/"
        f"{settings.embedding_monthly_call_limit} used; "
        f"{usage.remaining_calls} remaining."
    )


if __name__ == "__main__":
    main()
