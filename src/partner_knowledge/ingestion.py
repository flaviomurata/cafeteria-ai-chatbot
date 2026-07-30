"""Build the persistent local Partner knowledge index from approved sources."""

import csv
import hashlib
import json
from collections.abc import Callable, Iterable
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from uuid import uuid4
from zipfile import BadZipFile

import chromadb
from chromadb.errors import ChromaError
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from src.config import get_settings
from src.partner_knowledge.config import (
    PartnerKnowledgeSettings,
    get_partner_knowledge_settings,
)

_COLLECTION_NAME = "partner_knowledge"
_STAGING_COLLECTION_PREFIX = "partner_knowledge_staging_"
_APPROVED_SOURCES = {
    "Catálogo de Produtos e Ingredientes — Café Aurora.pdf": (
        "Catálogo de Produtos e Ingredientes — Café Aurora",
        "pdf",
    ),
    "CA-COM-PLA-001_Controle_de_Estoque.csv": ("Controle de Estoque", "csv"),
    "CA-TEC-CAD-001_Configuracao_das_Unidades.json": (
        "Configuração das Unidades",
        "json",
    ),
    "Manual de Operações das Unidades — Café Aurora.pdf": (
        "Manual de Operações das Unidades — Café Aurora",
        "pdf",
    ),
    "CA-QUA-GUI-001_Guia_de_Atendimento_ao_Cliente.docx": (
        "Guia de Atendimento ao Cliente",
        "docx",
    ),
    "CA-FIN-POL-001_Politica_de_Despesas_e_Reembolsos.docx": (
        "Política de Despesas e Reembolsos",
        "docx",
    ),
}


class PartnerKnowledgeIngestionError(RuntimeError):
    """Raised when approved Partner knowledge cannot be safely indexed."""


@dataclass(frozen=True)
class IngestionResult:
    indexed_chunks: int
    index_path: Path


@dataclass(frozen=True)
class _Chunk:
    text: str
    document_name: str
    location: str
    technical_location: str
    source_kind: str

    @property
    def id(self) -> str:
        identity = "\n".join((self.document_name, self.technical_location, self.text))
        return hashlib.sha256(identity.encode()).hexdigest()

    @property
    def metadata(self) -> dict[str, str]:
        return {
            "document_name": self.document_name,
            "location": self.location,
            "technical_location": self.technical_location,
            "source_kind": self.source_kind,
        }


class PartnerKnowledgeIngestor:
    """Explicit, idempotent builder for the approved Partner knowledge index."""

    def __init__(
        self,
        settings: PartnerKnowledgeSettings,
        *,
        embed_documents: Callable[[list[str]], list[list[float]]],
    ):
        self._settings = settings
        self._embed_documents = embed_documents

    def ingest(self) -> IngestionResult:
        chunks = list(self._load_approved_chunks())
        if not chunks:
            raise PartnerKnowledgeIngestionError(
                "No approved Partner knowledge documents were found in "
                f"{self._settings.partner_document_source}."
            )
        try:
            embeddings = self._embed_documents([chunk.text for chunk in chunks])
        except Exception as exc:
            raise PartnerKnowledgeIngestionError(
                "Unable to create embeddings for Partner knowledge."
            ) from exc
        if len(embeddings) != len(chunks):
            raise PartnerKnowledgeIngestionError(
                "Embedding provider returned a different number of vectors "
                "than documents."
            )
        self._settings.partner_index_path.mkdir(parents=True, exist_ok=True)
        staging_name = f"{_STAGING_COLLECTION_PREFIX}{uuid4().hex}"
        client = None
        try:
            client = chromadb.PersistentClient(
                path=str(self._settings.partner_index_path)
            )
            collection = client.create_collection(
                staging_name,
                embedding_function=None,
                metadata={"hnsw:space": "cosine"},
            )
            collection.add(
                ids=[chunk.id for chunk in chunks],
                documents=[chunk.text for chunk in chunks],
                metadatas=[chunk.metadata for chunk in chunks],
                embeddings=embeddings,
            )
        except (ChromaError, OSError, ValueError) as exc:
            if client is not None:
                try:
                    client.delete_collection(staging_name)
                except ChromaError:
                    pass
            raise PartnerKnowledgeIngestionError(
                "Unable to write the Partner knowledge index at "
                f"{self._settings.partner_index_path}."
            ) from exc
        try:
            try:
                client.delete_collection(_COLLECTION_NAME)
            except ChromaError:
                pass
            collection.modify(name=_COLLECTION_NAME)
        except ChromaError as exc:
            raise PartnerKnowledgeIngestionError(
                "Unable to activate the rebuilt Partner knowledge index at "
                f"{self._settings.partner_index_path}."
            ) from exc
        return IngestionResult(
            indexed_chunks=len(chunks), index_path=self._settings.partner_index_path
        )

    def _load_approved_chunks(self) -> Iterable[_Chunk]:
        source_root = self._settings.partner_document_source
        if not source_root.is_dir():
            raise PartnerKnowledgeIngestionError(
                f"Partner document source is not a directory: {source_root}"
            )
        for filename, (document_name, source_kind) in _APPROVED_SOURCES.items():
            path = source_root / filename
            if not path.is_file():
                continue
            if source_kind == "pdf":
                yield from self._extract_pdf(path, document_name)
            elif source_kind == "docx":
                yield from self._extract_docx(path, document_name)
            elif source_kind == "csv":
                yield from self._extract_csv(path, document_name)
            else:
                yield from self._extract_json(path, document_name)

    @staticmethod
    def _extract_pdf(path: Path, document_name: str) -> Iterable[_Chunk]:
        try:
            pages = PdfReader(path).pages
            extracted_pages = [page.extract_text().strip() for page in pages]
        except (PdfReadError, OSError) as exc:
            raise PartnerKnowledgeIngestionError(
                f"Approved document {path.name} has no readable native text; "
                "OCR is not supported."
            ) from exc
        if not any(extracted_pages):
            raise PartnerKnowledgeIngestionError(
                f"Approved document {path.name} has no readable native text; "
                "OCR is not supported."
            )
        for page_number, text in enumerate(extracted_pages, start=1):
            if text:
                yield _Chunk(
                    text,
                    document_name,
                    f"Página {page_number}",
                    f"page:{page_number}",
                    "pdf",
                )

    @staticmethod
    def _extract_docx(path: Path, document_name: str) -> Iterable[_Chunk]:
        try:
            document = Document(path)
        except (BadZipFile, PackageNotFoundError, OSError) as exc:
            raise PartnerKnowledgeIngestionError(
                f"Approved document {path.name} has no readable native text; "
                "OCR is not supported."
            ) from exc
        paragraphs = document.paragraphs
        sections: list[tuple[str, list[str]]] = []
        heading, body = "Introdução", []
        for paragraph in paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style.name.startswith("Heading"):
                if body:
                    sections.append((heading, body))
                heading, body = text, []
            else:
                body.append(text)
        if body:
            sections.append((heading, body))
        if not sections:
            raise PartnerKnowledgeIngestionError(
                f"Approved document {path.name} has no readable native text; "
                "OCR is not supported."
            )
        for index, (heading, body) in enumerate(sections, start=1):
            yield _Chunk(
                "\n".join([heading, *body]),
                document_name,
                f"Seção: {heading}",
                f"section:{index}",
                "docx",
            )
        for table_number, table in enumerate(document.tables, start=1):
            for row_number, row in enumerate(table.rows, start=1):
                cells = [cell.text.strip() for cell in row.cells]
                if not any(cells):
                    continue
                yield _Chunk(
                    " | ".join(cells),
                    document_name,
                    f"Tabela {table_number}, linha {row_number}",
                    f"table:{table_number}:row:{row_number}",
                    "docx",
                )

    @staticmethod
    def _extract_csv(path: Path, document_name: str) -> Iterable[_Chunk]:
        try:
            with path.open(encoding="utf-8-sig", newline="") as source_file:
                for row_number, row in enumerate(csv.DictReader(source_file), start=2):
                    values = {key: value for key, value in row.items() if value}
                    if not values:
                        continue
                    unit = values.get("Código da unidade", "unidade não identificada")
                    item = values.get("Código do item", "item não identificado")
                    description = values.get("Descrição", "")
                    location = f"Unidade {unit}, item {item}" + (
                        f" — {description}" if description else ""
                    )
                    yield _Chunk(
                        "\n".join(f"{key}: {value}" for key, value in values.items()),
                        document_name,
                        location,
                        f"row:{row_number}",
                        "csv",
                    )
        except (OSError, UnicodeDecodeError, csv.Error) as exc:
            raise PartnerKnowledgeIngestionError(
                f"Approved document {path.name} could not be read as CSV."
            ) from exc

    @staticmethod
    def _extract_json(path: Path, document_name: str) -> Iterable[_Chunk]:
        try:
            content: dict[str, Any] = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, UnicodeDecodeError, json.JSONDecodeError) as exc:
            raise PartnerKnowledgeIngestionError(
                f"Approved document {path.name} could not be read as JSON."
            ) from exc
        global_configuration = {
            key: value
            for key, value in content.items()
            if key not in {"unidades", "documento"}
        }
        if global_configuration:
            yield _Chunk(
                json.dumps(global_configuration, ensure_ascii=False, indent=2),
                document_name,
                "Configurações globais",
                "json:global",
                "json",
            )
        for unit_index, unit in enumerate(content.get("unidades", [])):
            unit_code, unit_name = (
                unit.get("codigo", "unidade não identificada"),
                unit.get("nome", ""),
            )
            location = f"Unidade {unit_code}" + (f" — {unit_name}" if unit_name else "")
            yield _Chunk(
                json.dumps(unit, ensure_ascii=False, indent=2),
                document_name,
                location,
                f"json:unidades[{unit_index}]",
                "json",
            )


def main() -> None:
    """Run explicit local ingestion with Gemini document embeddings."""
    from langchain_google_genai import GoogleGenerativeAIEmbeddings

    settings = get_partner_knowledge_settings()
    embeddings = GoogleGenerativeAIEmbeddings(
        model=settings.embedding_model, google_api_key=get_settings().google_api_key
    )
    result = PartnerKnowledgeIngestor(
        settings, embed_documents=embeddings.embed_documents
    ).ingest()
    print(
        "Indexed "
        f"{result.indexed_chunks} Partner knowledge chunks at {result.index_path}"
    )


if __name__ == "__main__":
    main()
